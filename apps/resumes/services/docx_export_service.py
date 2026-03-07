"""
DOCX Export Service for generating Word document resumes.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.shortcuts import get_object_or_404
from ..models import Resume
import logging
import io

logger = logging.getLogger(__name__)


class DOCXExportService:
    """
    Service class for generating DOCX exports of resumes.
    Uses python-docx to create formatted Word documents.
    """
    
    @staticmethod
    def _hex_to_rgb(hex_color):
        """
        Convert hex color to RGB tuple.
        
        Args:
            hex_color: Hex color string (e.g., '#2C3E50')
            
        Returns:
            tuple: (r, g, b) values
        """
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

    @staticmethod
    def generate_docx(resume_id, version_id=None):
        """
        Generate DOCX from resume and return as bytes.
        Applies custom fonts and colors from resume settings.
        
        Args:
            resume_id: ID of the resume to export
            version_id: Optional ID of specific version to export
        
        Returns:
            tuple: (bytes, Resume) - DOCX file content as bytes and Resume object
        
        Raises:
            Resume.DoesNotExist: If resume with given ID doesn't exist
        """
        try:
            # Load resume with all related sections
            resume = get_object_or_404(
                Resume.objects.prefetch_related(
                    'personal_info',
                    'experiences',
                    'education',
                    'skills',
                    'projects'
                ),
                id=resume_id
            )
            
            # If version_id is provided, create a temporary resume from version snapshot
            if version_id:
                from ..models import ResumeVersion
                from .snapshot_utils import create_resume_from_snapshot
                version = get_object_or_404(ResumeVersion, id=version_id, resume=resume)
                resume = create_resume_from_snapshot(resume, version.snapshot_data)
            
            # Get customization settings (Requirements: 14.5)
            from .template_customization_service import TemplateCustomizationService
            color_scheme = TemplateCustomizationService.get_color_scheme(resume.color_scheme)
            font_info = TemplateCustomizationService.get_font_family(resume.font_family)
            
            # Create document
            document = Document()
            
            # Set default font for the entire document
            style = document.styles['Normal']
            font = style.font
            font.name = font_info['name']
            font.size = Pt(10)
            
            # Set document margins
            sections = document.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Add personal information
            DOCXExportService._add_personal_info(document, resume, color_scheme, font_info)
            
            # Add experience section
            if resume.experiences.exists():
                DOCXExportService._add_experiences(document, resume, color_scheme, font_info)
            
            # Add education section
            if resume.education.exists():
                DOCXExportService._add_education(document, resume, color_scheme, font_info)
            
            # Add skills section
            if resume.skills.exists():
                DOCXExportService._add_skills(document, resume, color_scheme, font_info)
            
            # Add projects section
            if resume.projects.exists():
                DOCXExportService._add_projects(document, resume, color_scheme, font_info)
            
            # Save to bytes
            docx_bytes = io.BytesIO()
            document.save(docx_bytes)
            docx_bytes.seek(0)
            
            logger.info(f'Successfully generated DOCX for resume {resume_id}' +
                       (f' version {version_id}' if version_id else ''))
            return docx_bytes.getvalue(), resume
            
        except Exception as e:
            logger.error(f'Failed to generate DOCX for resume {resume_id}: {str(e)}', exc_info=True)
            raise

    @staticmethod
    def _add_personal_info(document, resume, color_scheme, font_info):
        """Add personal information section to document with custom colors."""
        personal_info = getattr(resume, 'personal_info', None)
        if not personal_info:
            return
        
        # Convert primary color to RGB
        primary_rgb = DOCXExportService._hex_to_rgb(color_scheme['primary'])
        secondary_rgb = DOCXExportService._hex_to_rgb(color_scheme['secondary'])
        
        # Name (centered, large, bold, custom color)
        name_paragraph = document.add_paragraph()
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_paragraph.add_run(personal_info.full_name)
        name_run.font.name = font_info['name']
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(*primary_rgb)
        
        # Contact information (centered)
        contact_paragraph = document.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_parts = []
        
        if personal_info.phone:
            contact_parts.append(personal_info.phone)
        if personal_info.email:
            contact_parts.append(personal_info.email)
        if personal_info.location:
            contact_parts.append(personal_info.location)
        
        contact_run = contact_paragraph.add_run(' | '.join(contact_parts))
        contact_run.font.name = font_info['name']
        contact_run.font.size = Pt(10)
        
        # Links (centered, custom color)
        if personal_info.linkedin or personal_info.github:
            links_paragraph = document.add_paragraph()
            links_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            links_parts = []
            
            if personal_info.linkedin:
                links_parts.append(personal_info.linkedin)
            if personal_info.github:
                links_parts.append(personal_info.github)
            
            links_run = links_paragraph.add_run(' | '.join(links_parts))
            links_run.font.name = font_info['name']
            links_run.font.size = Pt(9)
            links_run.font.color.rgb = RGBColor(*secondary_rgb)
        
        # Add spacing
        document.add_paragraph()

    @staticmethod
    def _add_section_heading(document, heading_text, color_scheme, font_info):
        """Add a section heading with formatting and custom colors."""
        primary_rgb = DOCXExportService._hex_to_rgb(color_scheme['primary'])
        
        heading = document.add_heading(heading_text, level=1)
        heading_run = heading.runs[0]
        heading_run.font.name = font_info['name']
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(*primary_rgb)
        
        # Add horizontal line
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6)

    @staticmethod
    def _add_experiences(document, resume, color_scheme, font_info):
        """Add work experience section to document with custom styling."""
        DOCXExportService._add_section_heading(document, 'WORK EXPERIENCE', color_scheme, font_info)
        
        primary_rgb = DOCXExportService._hex_to_rgb(color_scheme['primary'])
        
        for experience in resume.experiences.all():
            # Company and role (bold, custom color)
            title_paragraph = document.add_paragraph()
            title_run = title_paragraph.add_run(f"{experience.role} | {experience.company}")
            title_run.font.name = font_info['name']
            title_run.font.bold = True
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = RGBColor(*primary_rgb)
            
            # Dates
            date_paragraph = document.add_paragraph()
            date_paragraph.paragraph_format.space_before = Pt(0)
            date_paragraph.paragraph_format.space_after = Pt(3)
            
            start_date = experience.start_date.strftime('%B %Y')
            end_date = experience.end_date.strftime('%B %Y') if experience.end_date else 'Present'
            date_run = date_paragraph.add_run(f"{start_date} - {end_date}")
            date_run.font.name = font_info['name']
            date_run.font.size = Pt(10)
            date_run.font.italic = True
            
            # Description (bullet points)
            if experience.description:
                # Split description into bullet points
                bullets = [line.strip() for line in experience.description.split('\n') if line.strip()]
                for bullet in bullets:
                    # Remove existing bullet markers
                    bullet_text = bullet.lstrip('•-* ')
                    bullet_paragraph = document.add_paragraph(bullet_text, style='List Bullet')
                    bullet_paragraph.paragraph_format.left_indent = Inches(0.25)
                    bullet_paragraph.paragraph_format.space_after = Pt(3)
                    for run in bullet_paragraph.runs:
                        run.font.name = font_info['name']
                        run.font.size = Pt(10)
            
            # Add spacing between experiences
            document.add_paragraph()

    @staticmethod
    def _add_education(document, resume, color_scheme, font_info):
        """Add education section to document with custom styling."""
        DOCXExportService._add_section_heading(document, 'EDUCATION', color_scheme, font_info)
        
        primary_rgb = DOCXExportService._hex_to_rgb(color_scheme['primary'])
        
        for edu in resume.education.all():
            # Degree and field (bold, custom color)
            title_paragraph = document.add_paragraph()
            title_run = title_paragraph.add_run(f"{edu.degree} in {edu.field}")
            title_run.font.name = font_info['name']
            title_run.font.bold = True
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = RGBColor(*primary_rgb)
            
            # Institution and years
            details_paragraph = document.add_paragraph()
            details_paragraph.paragraph_format.space_before = Pt(0)
            details_paragraph.paragraph_format.space_after = Pt(6)
            
            end_year = edu.end_year if edu.end_year else 'Present'
            details_run = details_paragraph.add_run(f"{edu.institution} | {edu.start_year} - {end_year}")
            details_run.font.name = font_info['name']
            details_run.font.size = Pt(10)
            details_run.font.italic = True
            
            # Add spacing between education entries
            document.add_paragraph()

    @staticmethod
    def _add_skills(document, resume, color_scheme, font_info):
        """Add skills section to document with custom styling."""
        DOCXExportService._add_section_heading(document, 'SKILLS', color_scheme, font_info)
        
        primary_rgb = DOCXExportService._hex_to_rgb(color_scheme['primary'])
        
        # Group skills by category
        skills_by_category = {}
        for skill in resume.skills.all():
            category = skill.category or 'General'
            if category not in skills_by_category:
                skills_by_category[category] = []
            skills_by_category[category].append(skill.name)
        
        # Add each category
        for category, skill_names in skills_by_category.items():
            paragraph = document.add_paragraph()
            
            # Category name (bold, custom color)
            category_run = paragraph.add_run(f"{category}: ")
            category_run.font.name = font_info['name']
            category_run.font.bold = True
            category_run.font.size = Pt(10)
            category_run.font.color.rgb = RGBColor(*primary_rgb)
            
            # Skills (comma-separated)
            skills_run = paragraph.add_run(', '.join(skill_names))
            skills_run.font.name = font_info['name']
            skills_run.font.size = Pt(10)
            
            paragraph.paragraph_format.space_after = Pt(3)
        
        # Add spacing
        document.add_paragraph()

    @staticmethod
    def _add_projects(document, resume, color_scheme, font_info):
        """Add projects section to document with custom styling."""
        DOCXExportService._add_section_heading(document, 'PROJECTS', color_scheme, font_info)
        
        primary_rgb = DOCXExportService._hex_to_rgb(color_scheme['primary'])
        secondary_rgb = DOCXExportService._hex_to_rgb(color_scheme['secondary'])
        
        for project in resume.projects.all():
            # Project name (bold, custom color)
            title_paragraph = document.add_paragraph()
            title_run = title_paragraph.add_run(project.name)
            title_run.font.name = font_info['name']
            title_run.font.bold = True
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = RGBColor(*primary_rgb)
            
            # URL if available
            if project.url:
                title_run.add_text(' | ')
                url_run = title_paragraph.add_run(project.url)
                url_run.font.name = font_info['name']
                url_run.font.size = Pt(9)
                url_run.font.color.rgb = RGBColor(*secondary_rgb)
            
            # Description
            if project.description:
                desc_paragraph = document.add_paragraph(project.description)
                desc_paragraph.paragraph_format.space_before = Pt(0)
                desc_paragraph.paragraph_format.space_after = Pt(3)
                for run in desc_paragraph.runs:
                    run.font.name = font_info['name']
                    run.font.size = Pt(10)
            
            # Technologies
            if project.technologies:
                tech_paragraph = document.add_paragraph()
                tech_paragraph.paragraph_format.space_before = Pt(0)
                tech_paragraph.paragraph_format.space_after = Pt(6)
                
                tech_label = tech_paragraph.add_run('Technologies: ')
                tech_label.font.name = font_info['name']
                tech_label.font.bold = True
                tech_label.font.size = Pt(9)
                tech_label.font.italic = True
                
                tech_run = tech_paragraph.add_run(project.technologies)
                tech_run.font.name = font_info['name']
                tech_run.font.size = Pt(9)
                tech_run.font.italic = True
            
            # Add spacing between projects
            document.add_paragraph()
